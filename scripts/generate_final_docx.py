import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

def set_font(run, size, name='Times New Roman'):
    run.font.name = name
    run.font.size = Pt(size)
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

def add_styled_paragraph(doc, text, size, bold=False, heading=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if heading:
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(6)
    else:
        # Adjust line spacing to fill pages
        p.paragraph_format.line_spacing = 1.5 
        p.paragraph_format.space_after = Pt(10)
        
    run = p.add_run(text)
    set_font(run, size)
    run.bold = bold
    return p

def generate_report():
    input_path = r"d:\Codes\DIP Project\reports\Project_Report_Final.md"
    output_path = r"d:\Codes\DIP Project\reports\FINAL_PROJECT_REPORT.docx"
    
    doc = Document()
    
    # Page setup for 16 pages
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Handle markers for screenshots
        if "[Screenshot" in line:
            doc.add_page_break()
            p = add_styled_paragraph(doc, line.replace("[", "").replace("]", ""), 14, bold=True, heading=True)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Insert demo images if they exist
            demo_dir = r"d:\Codes\DIP Project\outputs\demo"
            if "Dashboard" in line and os.path.exists(os.path.join(demo_dir, "timeline_demo.png")):
                doc.add_picture(os.path.join(demo_dir, "timeline_demo.png"), width=Inches(6))
                add_styled_paragraph(doc, "Figure 1: Real-time detection timeline and confidence tracking.", 10, bold=False)
            elif "Heatmap" in line and os.path.exists(os.path.join(demo_dir, "confusion_demo.png")):
                doc.add_picture(os.path.join(demo_dir, "confusion_demo.png"), width=Inches(5))
                add_styled_paragraph(doc, "Figure 2: Confusion Matrix representing model classification accuracy.", 10, bold=False)
            continue

        # Handle headings
        if line.startswith("# "):
            add_styled_paragraph(doc, line[2:].upper(), 18, bold=True, heading=True)
        elif line.startswith("## "):
            add_styled_paragraph(doc, line[3:], 14, bold=True, heading=True)
        elif line.startswith("### "):
            add_styled_paragraph(doc, line[4:], 13, bold=True, heading=True)
        elif line.startswith("---"):
            doc.add_page_break()
        else:
            # Handle list items
            if line.startswith("- ") or line.startswith("* "):
                add_styled_paragraph(doc, "  • " + line[2:], 12)
            elif re.match(r'^\d+\.', line):
                add_styled_paragraph(doc, line, 12)
            else:
                add_styled_paragraph(doc, line, 12)

    # Final result check
    doc.save(output_path)
    print(f"Report generated successfully: {output_path}")

if __name__ == "__main__":
    generate_report()
