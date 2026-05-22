import os
import re
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_font(run, size, name='Times New Roman', bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), name)

def set_two_columns(section):
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')
    if not cols:
        cols = OxmlElement('w:cols')
        sectPr.append(cols)
    else:
        cols = cols[0]
    cols.set(qn('w:num'), '2')
    cols.set(qn('w:space'), '280') 

def add_ieee_heading(doc, text, level):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if level == 1 else WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    
    if level == 1:
        run = p.add_run(text.upper())
        set_font(run, 10, bold=True)
    elif level == 2:
        run = p.add_run(text)
        set_font(run, 10, italic=True)
    return p

def generate_full_dissertation():
    input_path = r"d:\Codes\DIP Project\reports\Project_Report_Final.md"
    assets_dir = r"d:\Codes\DIP Project\assets"
    output_path = r"d:\Codes\DIP Project\reports\Full_Research_Paper_v4_HyperDense.docx"
    
    doc = Document()
    
    # Page Margins
    section = doc.sections[0]
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    
    # Title Block (Single Column)
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_p.add_run("Hybrid-TCN: A Multi-Stream Temporal Convolutional Framework for Real-Time Micro-Expression Recognition")
    set_font(title_run, 24, bold=True)
    
    auth_p = doc.add_paragraph()
    auth_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    auth_run = auth_p.add_run("P. Prateek, Y. Srinadh, Abhishek Das, G. Sai Tejesh, K. Hanok\n")
    set_font(auth_run, 12)
    dept_run = auth_p.add_run("School of Engineering and Sciences, SRM University – AP, Andhra Pradesh\n")
    set_font(dept_run, 11, italic=True)

    # Section Break for Two Columns
    doc.add_section()
    new_section = doc.sections[-1]
    set_two_columns(new_section)

    with open(input_path, 'r', encoding='utf-8') as f:
        content = f.read()

    # Split into sections based on '##'
    blocks = re.split(r'\n## ', content)
    
    img_map = {
        "Architectural": "architecture.png",
        "Augmentation": "augmentation.png",
        "Receptive": "tcn_dilation.png",
        "Dashboard": "results.png",
        "Performance": "roc_curve.png"
    }

    for idx, block in enumerate(blocks):
        block = block.strip()
        if not block: continue
        
        # Skip the very first part if it's the title
        if idx == 0 and "HYBRID-TCN" in block:
            continue

        lines = block.split('\n')
        title = lines[0].strip()
        body = lines[1:]

        # Handle Abstract
        if "ABSTRACT" in title.upper():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            run = p.add_run("Abstract—")
            set_font(run, 10, bold=True, italic=True)
            run = p.add_run(" ".join(body).strip())
            set_font(run, 10, bold=True, italic=True)
            continue

        # Regular Sections
        if title:
            # Main section
            add_ieee_heading(doc, re.sub(r'^\d+\.', '', title).strip(), 1)

        for line in body:
            line = line.strip()
            if not line: continue
            
            # Sub-headings
            if line.startswith("### "):
                text = line.replace("### ", "").strip()
                # Check if it's a screenshot
                if "[Screenshot" in text:
                    # IMAGE LOGIC
                    best_match = None
                    for key, val in img_map.items():
                        if key.lower() in text.lower():
                            best_match = val
                            break
                    
                    if best_match:
                        img_path = os.path.join(assets_dir, best_match)
                        if os.path.exists(img_path):
                            p_img = doc.add_paragraph()
                            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_img = p_img.add_run()
                            run_img.add_picture(img_path, width=Inches(3.3))
                            
                            p_cap = doc.add_paragraph()
                            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            run_cap = p_cap.add_run(text.strip("[]"))
                            set_font(run_cap, 8, italic=True)
                        else:
                            print(f"Warning: Image missing at {img_path}")
                    continue
                else:
                    # Normal sub-heading
                    add_ieee_heading(doc, text, 2)
                    continue

            # Math
            if "$" in line:
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.replace("$", "").strip())
                set_font(run, 11, italic=True)
                continue

            # Tables (heuristic)
            if "|" in line:
                if "---" in line: continue
                p = doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(line.replace("|", "  ").strip())
                set_font(run, 9)
                continue

            # Paragraphs
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            p.paragraph_format.first_line_indent = Inches(0.14)
            p.paragraph_format.line_spacing = 1.15
            run = p.add_run(line.replace("**", ""))
            set_font(run, 10)

    doc.save(output_path)
    print(f"Final 10-page Dissertation generated: {output_path}")

if __name__ == "__main__":
    generate_full_dissertation()
